"""
Document Processing Service
Handles file uploads, text extraction, and coordinates the AI extraction pipeline
Optimized for Azure App Service deployment without GPU requirements
"""

import asyncio
import io
import tempfile
import time
import os
from pathlib import Path
from typing import Dict, List, Optional, Union, Any

import aiofiles
import structlog
from fastapi import UploadFile
from PIL import Image
import PyPDF2
import pypdf
from docx import Document as DocxDocument

# File type detection
try:
    import magic  # type: ignore
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

try:
    import filetype
    FILETYPE_AVAILABLE = True
except ImportError:
    FILETYPE_AVAILABLE = False

from ..config import get_settings
from ..models import ExtractResponse, TemplateExtractionResult, AIExtractionResult, ErrorResult
from .openrouter_client import OpenRouterClient


logger = structlog.get_logger(__name__)
settings = get_settings()

class DocumentProcessor:
    """Service for processing and extracting text from various document types."""
    
    def __init__(
        self, 
        openrouter_client: OpenRouterClient
    ):
        self.openrouter_client = openrouter_client
        self.magic_available = MAGIC_AVAILABLE
        self.filetype_available = FILETYPE_AVAILABLE
        
        if MAGIC_AVAILABLE:
            logger.info("Python-magic file type detection available")
        elif FILETYPE_AVAILABLE:
            logger.info("Filetype file type detection available")
        else:
            logger.warning("No file type detection libraries available - using extension-based detection only")
        
        # Define supported file types and their processors
        self.supported_types = {
            'application/pdf': self._extract_text_from_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_text_from_docx,
            'text/plain': self._extract_text_from_txt,
            'image/png': self._extract_text_from_image,
            'image/jpeg': self._extract_text_from_image,
            'image/jpg': self._extract_text_from_image,
        }
        
        # Processing statistics
        self.processing_stats = {
            "documents_processed": 0,
            "successful_extractions": 0,
            "failed_extractions": 0,
            "total_processing_time": 0.0,
            "average_processing_time": 0.0,
        }
    
    async def process_document(
        self,
        file: UploadFile,
        template: Optional[str] = None,
        custom_prompt: Optional[str] = None
    ) -> ExtractResponse:
        """Process uploaded document and extract information using AI"""
        start_time = time.time()
        
        try:
            logger.info("Starting document processing", filename=file.filename, template=template)
            
            # Validate file
            validation_result = await self.validate_document(file)
            if not validation_result["valid"]:
                return ErrorResult(
                    error=f"File validation failed: {', '.join(validation_result['errors'])}"
                )
            
            # Extract text from document
            try:
                text_content = await self._extract_text_from_file(file)
                if not text_content or not text_content.strip():
                    return ErrorResult(
                        error="No text content found - the document appears to be empty or contains no extractable text"
                    )
            except Exception as e:
                logger.error("Text extraction failed", error=str(e), filename=file.filename)
                return ErrorResult(
                    error=f"Text extraction failed: {str(e)}"
                )
            
            # Note: Documents are processed in-memory only - no permanent storage needed
            
            # Extract information using AI - Single-stage approach
            try:
                # Use single-stage extraction to eliminate JSON parsing issues
                result = await self.openrouter_client.extract_data_single_stage(
                    text_content, template, custom_prompt
                )
                
                # Update statistics
                processing_time = time.time() - start_time
                self._update_processing_stats(processing_time, True)
                
                logger.info("Document processing completed successfully", 
                           processing_time=processing_time)
                
                return result
                
            except Exception as e:
                logger.error("AI extraction failed", error=str(e), filename=file.filename)
                return ErrorResult(
                    error=f"AI extraction failed: {str(e)}"
                )
                
        except Exception as e:
            processing_time = time.time() - start_time
            self._update_processing_stats(processing_time, False)
            
            logger.error("Document processing failed", error=str(e), filename=file.filename)
            return ErrorResult(
                error=f"Document processing failed: {str(e)}"
            )
    
    async def extract_text_from_file(self, file: UploadFile) -> str:
        """
        Public method to extract text content from uploaded file
        Used by dynamic query service and other external callers
        """
        return await self._extract_text_from_file(file)
    
    async def _extract_text_from_file(self, file: UploadFile) -> str:
        """Extract text content from uploaded file"""
        # Read file content
        content = await file.read()
        
        # Reset file position for potential re-reading
        await file.seek(0)
        
        # Detect file type
        mime_type = self.detect_file_type(content, file.filename)
        
        # Extract text using appropriate processor
        if mime_type in self.supported_types:
            extractor = self.supported_types[mime_type]
            return await extractor(content)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
    
    def detect_file_type(self, file_content: bytes, filename: Optional[str] = None) -> str:
        """Detect file type using multiple methods for maximum reliability"""
        # Method 1: Try Filetype (Python 3.13 compatible detection)
        if self.filetype_available:
            try:
                kind = filetype.guess(file_content)
                if kind and kind.mime:
                    logger.debug(f"Filetype detected: {kind.mime}")
                    return kind.mime
            except Exception as e:
                logger.warning(f"Filetype detection failed: {e}")
        
        # Method 2: Extension-based detection as fallback
        if filename:
            extension = Path(filename).suffix.lower()
            extension_mapping = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.txt': 'text/plain',
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
            }
            
            if extension in extension_mapping:
                detected_type = extension_mapping[extension]
                logger.debug(f"Extension-based detection: {detected_type}")
                return detected_type
        
        # Method 3: Content-based detection as last resort
        detected_type = self._detect_by_content(file_content)
        if detected_type:
            logger.debug(f"Content-based detection: {detected_type}")
            return detected_type
        
        # Default fallback
        logger.warning("Could not determine file type, defaulting to text/plain")
        return 'text/plain'
    
    def _detect_by_content(self, content: bytes) -> Optional[str]:
        """Detect file type by examining file content signatures"""
        if not content or len(content) < 4:
            return None
        
        # Check file signatures (magic bytes)
        if content.startswith(b'%PDF'):
            return 'application/pdf'
        elif content.startswith(b'PK\x03\x04') and b'word/' in content[:1024]:
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif content.startswith(b'\x89PNG'):
            return 'image/png'
        elif content.startswith(b'\xff\xd8\xff'):
            return 'image/jpeg'
        
        # Try to decode as text
        try:
            content.decode('utf-8')
            return 'text/plain'
        except UnicodeDecodeError:
            pass
        
        return None
    
    async def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            # Create a BytesIO object from the file content
            pdf_file = io.BytesIO(file_content)
            
            # Try PyPDF2 first
            try:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                
                if text_content:
                    return "\n\n".join(text_content)
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}, trying pypdf")
            
            # Fallback to pypdf
            pdf_file.seek(0)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            if not text_content:
                raise ValueError("No text content found in PDF")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"PDF text extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_path)
                text_content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    if table_text:
                        text_content.append("\n".join(table_text))
                
                if not text_content:
                    raise ValueError("No text content found in DOCX")
                
                return "\n\n".join(text_content)
                
            finally:
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"DOCX text extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    if text.strip():
                        return text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"TXT text extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from TXT file: {str(e)}")
    
    async def _extract_text_from_image(self, file_content: bytes) -> str:
        """Extract text from image files - placeholder implementation"""
        try:
            image = Image.open(io.BytesIO(file_content))
            
            info = f"Image Processing Placeholder\n"
            info += f"Format: {image.format}\n"
            info += f"Size: {image.size[0]}x{image.size[1]} pixels\n"
            info += f"Mode: {image.mode}\n\n"
            info += "Note: OCR text extraction is not implemented in this version. "
            info += "To add OCR capabilities, consider integrating libraries like pytesseract or Azure Computer Vision."
            
            return info
            
        except Exception as e:
            logger.error(f"Image processing failed", error=str(e))
            raise ValueError(f"Failed to process image: {str(e)}")

    def _is_supported_file_type(self, filename: Optional[str]) -> bool:
        """Check if the file type is supported"""
        if not filename:
            return False
        
        extension = Path(filename).suffix.lower()
        return extension in ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg']
    
    def _update_processing_stats(self, processing_time: float, success: bool) -> None:
        """Update processing statistics"""
        self.processing_stats["documents_processed"] += 1
        self.processing_stats["total_processing_time"] += processing_time
        
        if success:
            self.processing_stats["successful_extractions"] += 1
        else:
            self.processing_stats["failed_extractions"] += 1
        
        # Calculate average processing time
        if self.processing_stats["documents_processed"] > 0:
            self.processing_stats["average_processing_time"] = (
                self.processing_stats["total_processing_time"] / 
                self.processing_stats["documents_processed"]
            )
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get current processing statistics"""
        return self.processing_stats.copy()
    
    async def get_available_templates(self) -> Dict[str, str]:
        """Get available extraction templates"""
        return settings.default_templates
    
    async def validate_document(self, file: UploadFile) -> Dict[str, Any]:
        """Validate uploaded document"""
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "file_info": {
                "filename": file.filename,
                "content_type": file.content_type,
                "size": None
            }
        }
        
        try:
            # Check file size
            if hasattr(file.file, 'seek') and hasattr(file.file, 'tell'):
                current_pos = file.file.tell()
                file.file.seek(0, 2)
                file_size = file.file.tell()
                file.file.seek(current_pos)
                
                validation_result["file_info"]["size"] = file_size
                
                max_size_bytes = settings.max_file_size_mb * 1024 * 1024
                if file_size > max_size_bytes:
                    validation_result["valid"] = False
                    validation_result["errors"].append(f"File size ({file_size} bytes) exceeds maximum allowed size ({max_size_bytes} bytes)")
            
            # Check file type
            if not self._is_supported_file_type(file.filename):
                validation_result["valid"] = False
                validation_result["errors"].append(f"Unsupported file type. Supported types: {settings.supported_file_types}")
            
            return validation_result
            
        except Exception as e:
            validation_result["valid"] = False
            validation_result["errors"].append(f"Validation error: {str(e)}")
            return validation_result

    def validate_file_size(self, file_size: int) -> bool:
        """Validate if file size is within allowed limits"""
        max_size = settings.max_file_size_mb * 1024 * 1024
        return file_size <= max_size
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported MIME types"""
        return list(self.supported_types.keys()) 